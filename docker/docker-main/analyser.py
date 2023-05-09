import os
import datetime
import docx
import bwa

class ExtraInfo:

    def __init__(self, lastPrinted, brutto, vat, netto, alreadyPaid=False):

        self.lastPrinted = lastPrinted
        self.brutto = brutto
        self.vat = vat
        self.netto = netto
        self.alreadyPaid = alreadyPaid

def analyseDocument(fileItem):

    if not fileItem.fileType == "Dokument (Rechnung)":
        return None

    TABLE_BANKING_INFO = 0
    TABLE_INVOICE_INFO = 0

    COL_STRING_INFO = 1
    COL_CURRENCY    = 2
    COL_VALUE       = 3
  
    invoice = None
    try:
        invoice = docx.Document(fileItem.fullpath)
    except docx.opc.exceptions.PackageNotFoundError as e:
        print("Error opening Invoice")
        print(e)
        return None

    table = invoice.tables[TABLE_INVOICE_INFO]
   
    netto  = 0
    vat    = 0
    brutto = 0
    try:
        netto  = float(table.rows[-3].cells[COL_VALUE].text.replace(".", "").replace(",","."))
        vat    = float(table.rows[-2].cells[COL_VALUE].text.replace(".", "").replace(",","."))
        brutto = float(table.rows[-1].cells[COL_VALUE].text.replace(".", "").replace(",","."))
    except IndexError:
        return None
    
    lastPrinted = invoice.core_properties.last_printed
    
    # check if ever printed #
    if lastPrinted < datetime.datetime(2001, 1, 1):
        lastPrinted = None
    
    bwaPaidState = bwa.getPaidStateForFile(fileItem.name)
    return ExtraInfo(lastPrinted, brutto, vat, netto, bwaPaidState) 
